import os
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # Import the namespace function if needed
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from openpyxl.drawing.image import Image as OpenpyxlImage
import matplotlib.pyplot as plt
from datetime import datetime
import pandas as pd
import matplotlib.image as mpimg

class ComplianceReportDocumentGenerator:
    def __init__(self, excel_file, client_name, services_link, logo_path=None):
        self.excel_file = excel_file
        self.client_name = client_name
        self.services_link = services_link
        self.logo_path = logo_path
        self.workbook = openpyxl.load_workbook(excel_file, read_only=False, data_only=True)
        self.document = Document()
        
        # Set custom page size (13x10 inches)
        sections = self.document.sections
        for section in sections:
            section.page_height = Inches(10)
            section.page_width = Inches(13)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # Add header with logo
            header = section.header
            header_table = header.add_table(1, 2, width=Inches(11))
            
            # Add page number to header
            header_cell_right = header_table.cell(0, 1)
            paragraph = header_cell_right.paragraphs[0]
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Add logo if provided
            if logo_path and os.path.exists(logo_path):
                header_cell_left = header_table.cell(0, 0)
                paragraph = header_cell_left.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _create_chart_from_excel_data(self, sheet_name, x_column, y_columns):
        """
        Create a matplotlib chart from Excel data
        
        Args:
            sheet_name (str): Name of the sheet
            x_column (str): Column to use for x-axis
            y_columns (list): Columns to use for y-axis
        
        Returns:
            str: Path to saved chart image
        """
        # Read sheet data
        sheet = self.workbook[sheet_name]
        data = list(sheet.values)
        headers = data[0]
        
        # Convert to DataFrame
        df = pd.DataFrame(data[1:], columns=headers)
        
        # Prepare chart
        plt.figure(figsize=(10, 6))
        
        # Plot based on sheet name
        if sheet_name == 'Priority Summary':
            # Priority distribution bar chart
            priorities = df['Priority'].tolist()
            counts = df['Count'].tolist()
            colors = ['green', 'orange', 'red', 'blue']
            plt.bar(priorities, counts, color=colors)
            plt.title('Priority Distribution')
            plt.xlabel('Priority')
            plt.ylabel('Count')
        
        elif sheet_name == 'Service Pivot':
            # Stacked bar chart for service priority distribution
            services = df['title'].tolist()
            high = df['High'].tolist()
            medium = df['Medium'].tolist()
            low = df['Low'].tolist()
            safe = df['Safe'].tolist()
            
            plt.figure(figsize=(15, 8))
            plt.bar(services, high, label='High', color='red')
            plt.bar(services, medium, bottom=high, label='Medium', color='orange')
            plt.bar(services, low, bottom=[h+m for h,m in zip(high, medium)], label='Low', color='yellow')
            plt.bar(services, safe, bottom=[h+m+l for h,m,l in zip(high, medium, low)], label='Safe', color='green')
            
            plt.title('Service Priority Distribution')
            plt.xlabel('Services')
            plt.ylabel('Count')
            plt.xticks(rotation=90)
            plt.legend()
            plt.tight_layout()
        
        # Save chart
        chart_path = f"{sheet_name.replace(' ', '_')}_chart.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return chart_path
    
    def _create_title_page(self):
        """Enhanced title page creation with better spacing"""
        title = self.document.add_heading(f"{self.client_name} AWS Report", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_paragraph = title.runs[0]
        title_paragraph.font.size = Pt(24)
        title_paragraph.font.color.rgb = RGBColor(31, 73, 125)
        
        # Add spacing
        self.document.add_paragraph().add_run().add_break()
        
        details = [
            f"Client: {self.client_name}",
            f"Report Name: {self.client_name} AWS Compliance Report",
            f"Report Date: {datetime.now().strftime('%A, %d %B %Y')}",
            "Report Version: Version 1.0"
        ]
        
        details_table = self.document.add_table(rows=len(details), cols=1)
        details_table.style = 'Table Grid'
        details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, detail in enumerate(details):
            cell = details_table.rows[i].cells[0]
            cell.text = detail
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing before index
        self.document.add_paragraph().add_run().add_break()
        self.document.add_paragraph().add_run().add_break()
        
        index_title = self.document.add_heading("Index", level=1)
        index_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Create enhanced index table with actual page numbers
        index_data = [
            ["Title", "Page"],
            ["Overview", "1"],
            ["Priority Summary", "2"],
            ["Service Pivot Analysis", "3"],
            ["Service Analysis", "4"],
            ["Link of Detailed Report", "5"],
            ["Synopsis", "6"],
            ["Conclusion", "7"]
        ]
        
        index_table = self.document.add_table(rows=len(index_data), cols=2)
        index_table.style = 'Table Grid'
        index_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row in enumerate(index_data):
            cells = index_table.rows[i].cells
            for j, cell_value in enumerate(row):
                cell = cells[j]
                cell.text = cell_value
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if i == 0:
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(12)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 73, 125)
        # Add page break after index
        self.document.add_page_break()

    def _set_column_widths(self, table, widths):
        """Set custom column widths for tables"""
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    

    def _add_table_to_doc(self, data, title=None):
        """Enhanced table formatting with better column widths and spacing"""
        if title:
            heading = self.document.add_heading(title, level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.runs[0].font.color.rgb = RGBColor(31, 73, 125)
            
            # Add spacing after heading
            self.document.add_paragraph().add_run().add_break()
        
        if not data:
            self.document.add_paragraph("No data available.")
            return
        
        table = self.document.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set custom column widths based on content type
        if "Service Analysis" in str(title):
            widths = [0.5, 1.0, 2.0, 5.0, 1.0, 1.0]  # Adjusted widths for Service Analysis
            self._set_column_widths(table, widths)
        
        color_map = {
            'High': RGBColor(231, 76, 60),
            'Medium': RGBColor(243, 156, 18),
            'Low': RGBColor(241, 196, 15),
            'Safe': RGBColor(46, 204, 113)
        }
        
        # Enhanced table formatting
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(cell_value) if cell_value is not None else ''
                paragraph = cell.paragraphs[0]
                
                # Center align all cells
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add spacing within cells
                paragraph.space_before = Pt(6)
                paragraph.space_after = Pt(6)
                
                if i == 0:
                    # Header formatting
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(11)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 73, 125)
                
                # Priority column formatting
                if ('priority' in str(data[0][j]).lower() or 
                    j == len(row)-1 and 'priority' in str(data[0]).lower()):
                    priority = str(cell_value).strip()
                    if priority in color_map:
                        run = paragraph.runs[0]
                        run.font.color.rgb = color_map[priority]
                        run.font.bold = True

        # Add spacing after table
        self.document.add_paragraph().add_run().add_break()
    
    def generate_comprehensive_report(self):
        """
        Generate a comprehensive Word document from the Excel report
        """
        self._create_title_page()
        
        # Add Overview section
        overview_paragraph = self.document.add_paragraph()
        run = overview_paragraph.add_run("Overview:")
        run.bold = True
        
        overview_text = (
            "The Cloud Security Posture Report is a comprehensive evaluation of your cloud environment, "
            "focusing on its alignment with critical industry standards like CIS, SOC 2, PCI DSS, and best practices. "
            "It provides actionable insights into potential risks and areas requiring attention, helping organizations maintain "
            "a robust security posture.\n\n"
            "Resource Inventory:\n"
            "• Offers a detailed count of all resources in your cloud environment, categorized by type (e.g., compute instances, "
            "storage buckets, managed databases).\n"
            "• Helps identify the breadth of resources being managed and ensures no critical resource is overlooked.\n\n"
            "Account/Region Breakdown:\n"
            "• Displays the distribution of resources across different accounts and regions.\n"
            "• Enables easy identification of resource concentration and highlights potential regional compliance concerns.\n\n"
            "Configuration Compliance:\n"
            "• Tracks key security and operational configurations, such as whether encryption, logging, or versioning is enabled.\n"
            "• Provides percentages to help assess how well your resources adhere to compliance and security best practices.\n\n"
            "Resource Age:\n"
            "• Identifies the age of resources, such as instances or snapshots, to assist in lifecycle management.\n"
            "• Helps in understanding resource utilization patterns and supports decisions on archiving, upgrading, or retiring resources."
        )
        self.document.add_paragraph(overview_text)


        # Sections to process
        sections = [
            {
                'name': 'Priority Summary',
                'type': 'mixed',
                'table_title': 'Priority Distribution'
            },
            {
                'name': 'Service Pivot',
                'type': 'mixed',
                'table_title': 'Service Priority Distribution'
            },
            {
                'name': 'Service Analysis',
                'type': 'table',
                'title': 'Service Category Analysis'
            }
        ]
        
        for section in sections:
            # Process each section
            self.document.add_page_break()
            
            # Add section title
            title = self.document.add_heading(section['name'], level=2)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add description for the table (if it is the 'Service Analysis' section)
            if section['name'] == 'Service Analysis':
                table_description = (
                    "This table provides an overview of critical compliance controls for Auto Scaling, EC2, S3, and more. "
                    "It includes descriptions, open issues, and priority levels to ensure security and operational efficiency. "
                    "Refer to the attached Excel sheet for details."
                )
                self.document.add_paragraph(table_description)
            
            # Extract and add data based on section type
            if section['type'] in ['table', 'mixed']:
                # Extract table data
                table_data = self._extract_table_data(section['name'])
                
                # Add table
                if table_data:
                    self._add_table_to_doc(
                        table_data, 
                        title=section.get('table_title', section['name'])
                    )
                
                # For mixed sections, also create chart
                if section['type'] == 'mixed':
                    try:
                        chart_path = self._create_chart_from_excel_data(
                            section['name'], 
                            'Priority' if section['name'] == 'Priority Summary' else 'title', 
                            ['Count'] if section['name'] == 'Priority Summary' else ['High', 'Medium', 'Low', 'Safe']
                        )
                        # Add description for the chart
                        chart_description = (
                            "This chart illustrates the allocation of control titles across different AWS services, offering a comprehensive view "
                            "of the implemented controls for compliance management. It emphasizes the service areas with the highest concentration "
                            "of controls, helping prioritize resource allocation effectively."
                        )
                        self.document.add_paragraph(chart_description)
                        
                        self.document.add_picture(chart_path, width=Inches(10))
                    except Exception as e:
                        self.document.add_paragraph(f"Could not create chart: {e}")
        
        # Add detailed report link section with Key Components in a table format
        self.document.add_page_break()
        link_section = self.document.add_heading("Link of Detailed Report", level=2)
        link_section.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add text explaining the detailed information available in the link
        link_description = (
            "In this link, you will find detailed information for each service, including its ID, ARN, and other essential details "
            "such as title, control_title, control_description, region, account_id, resource, reason, description, priority, "
            "Recommendation Steps/Approach, and status. For example, the following data will be available for each service:\n\n"
            
            "**Example 1:**\n"
            "**Service**: ACM\n"
            "**Control Title**: ACM certificates should not expire within 30 days\n"
            "**Control Description**: Ensure network integrity is protected by ensuring X509 certificates are issued by AWS ACM.\n"
            "**Region**: ap-south-1\n"
            "**Account ID**: [Account ID]\n"
            "**ARN**: [ARN Information]\n"
            "**Resource**: [Resource Expiry Information]\n"
            "**Reason**: This section contains recommendations for configuring ACM resources.\n"
            "**Priority**: High\n"
            "**Recommendation Steps/Approach**: Ensure the AWS account is a member of AWS Organizations. Steps: 1. Review account settings. 2. Enroll in AWS Organizations if not already a member.\n"
            "**Status**: ok\n\n"
            
            "**Detailed Explanation:**\n"
            "**Service**: Indicates that the compliance check pertains to AWS ACM (Certificate Manager).\n\n"
            "**Control Title**: The control ensures that no ACM certificates in use are set to expire within the next 30 days. "
            "This is critical for maintaining secure communication and avoiding service disruptions.\n\n"
            "**Control Description**: Emphasizes the importance of using X509 certificates issued by AWS ACM to ensure network integrity "
            "and secure connections.\n\n"
            "**Region**: Specifies the geographical region where this compliance check applies, which is ap-south-1 in this case.\n\n"
            "**Account ID**: Placeholder for the unique identifier of the AWS account being evaluated.\n\n"
            "**ARN**: Placeholder for the Amazon Resource Name, which uniquely identifies the specific resource being referred to.\n\n"
            "**Resource**: Placeholder for information regarding the expiration details of the certificates under review.\n\n"
            "**Reason**: Indicates the purpose of this compliance check, which is to provide recommendations for configuring ACM resources effectively.\n\n"
            "**Priority Levels:**\n"
            "- **High**: Requires immediate attention and is critical.\n"
            "- **Medium**: After resolving high-priority items, medium-priority issues can be addressed.\n"
            "- **Low**: Issues that are less urgent and can be addressed as time permits.\n"
            "- **Safe/Well-Architected**: These resources already follow best practices and meet all requirements; no action is needed.\n\n"
            "**Status:**\n"
            "- **ok**: The resource is compliant with best practices.\n"
            "- **info**: Additional information is provided, but no action is required.\n"
            "- **skip**: The check was skipped as it does not apply in this context.\n"
            "- **alarm**: Action is required to bring the resource into compliance with best practices.\n"
        )
        
        self.document.add_paragraph(link_description)
        
        # Create a table for Key Components
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Adding the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S.No"
        hdr_cells[1].text = "Category"
        
        # Adding the rows for each component
        components = [
            ("1", "Security Services"),
            ("2", "Compute Services"),
            ("3", "Storage Services"),
            ("4", "Network Services"),
            ("5", "Database Services"),
            ("6", "Safe and Unsafe Issues"),
            ("7", "Analysis Graph")
        ]
        
        for component in components:
            row_cells = table.add_row().cells
            row_cells[0].text = component[0]
            row_cells[1].text = component[1]
        
        # Add the detailed report link with hyperlink formatting
        link_para = self.document.add_paragraph()
        link_run = link_para.add_run("Link: ")
        link_run.bold = True
        
        hyperlink_run = link_para.add_run(self.services_link)
        hyperlink_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
        hyperlink_run.underline = True
        
        # Add synopsis section
        self.document.add_page_break()
        synopsis_section = self.document.add_heading("Synopsis", level=2)
        synopsis_section.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        synopsis_text = (
            "This report outlines the security and compliance status of AWS resources, focusing on best practices and recommendations "
            "provided by AWS Security Hub and other AWS services. The goal is to ensure that AWS environments are configured to minimize "
            "security risks, maintain compliance, and safeguard sensitive data and resources. The sections below cover different aspects of "
            "the AWS security landscape, including Identity and Access Management (IAM), Key Management Service (KMS), AWS GuardDuty, "
            "CloudWatch, CloudTrail, and more. Each section will detail specific findings, associated risks, and compliance levels, as well "
            "as actionable recommendations for enhancing the overall security posture.\n\n"
            
            "1. IAM (Identity and Access Management):\n"
            "- Description: IAM is crucial for managing access to AWS services and resources securely. It allows fine-grained access control "
            "and ensures that users and roles have only the necessary permissions.\n"
            "- Key Findings: Ensure IAM password policies expire passwords within 90 days or less. Prevent password reuse and require a minimum "
            "length of 14 characters. Ensure that policies are applied to groups or roles, not users directly. Enable MFA for all admin users.\n"
            "- Recommendations: Regularly review IAM policies to ensure password strength, role configurations, and MFA settings are enforced.\n\n"

            "2. KMS (Key Management Service):\n"
            "- Description: KMS helps with data encryption and key management, ensuring that data is kept secure both at rest and in transit.\n"
            "- Key Findings: Enable automatic key rotation for Customer Managed Keys (CMKs).\n"
            "- Recommendations: Implement key rotation to enhance data confidentiality and integrity.\n\n"

            "3. GuardDuty:\n"
            "- Description: GuardDuty is a threat detection service that continuously monitors AWS accounts and workloads for suspicious activity.\n"
            "- Key Findings: Monitor and analyze alerts generated by GuardDuty to identify potential security risks.\n"
            "- Recommendations: Ensure that GuardDuty is enabled across all regions and integrate with other security tools for better visibility.\n\n"

            "4. CloudWatch:\n"
            "- Description: CloudWatch enables monitoring of AWS resources and applications, providing real-time insights into system performance "
            "and security events.\n"
            "- Key Findings: Ensure CloudWatch alarms are set up for critical security events.\n"
            "- Recommendations: Leverage CloudWatch for automated monitoring and incident detection.\n\n"

            "5. CloudTrail:\n"
            "- Description: CloudTrail logs all API calls and resource changes, providing a valuable audit trail for compliance and security audits.\n"
            "- Key Findings: Ensure that CloudTrail trails are enabled in all regions. Configure trails to integrate with CloudWatch for centralized logging.\n"
            "- Recommendations: Review CloudTrail logs regularly for unusual activities or policy violations.\n\n"

            "6. VPC (Virtual Private Cloud):\n"
            "- Description: VPC enables you to create isolated networks within the AWS environment, providing secure communication between resources.\n"
            "- Key Findings: Restrict ingress to remote server administration ports (SSH, RDP) to avoid unauthorized access.\n"
            "- Recommendations: Regularly review security group configurations and ensure that traffic is limited to necessary sources.\n\n"

            "7. Secrets Manager:\n"
            "- Description: AWS Secrets Manager securely stores and manages sensitive information like database credentials and API keys.\n"
            "- Key Findings: Ensure that all sensitive information is encrypted and managed through Secrets Manager.\n"
            "- Recommendations: Use Secrets Manager to reduce the risks of hardcoded secrets and ensure secure access management.\n\n"

            "8. Backup:\n"
            "- Description: Backup services provide automated backup solutions to safeguard critical data against loss or corruption.\n"
            "- Key Findings: Ensure backup plans and vaults are configured across all regions.\n"
            "- Recommendations: Implement backup plans with redundancy across regions to ensure data durability and disaster recovery.\n\n"

            "9. EC2 (Elastic Compute Cloud):\n"
            "- Description: EC2 provides scalable computing capacity in the cloud, and securing EC2 instances is vital for maintaining overall system integrity.\n"
            "- Key Findings: Ensure EC2 instances have termination protection and are using IAM profiles for secure access.\n"
            "- Recommendations: Regularly monitor EC2 instances for security vulnerabilities and ensure secure configurations are applied.\n\n"

            "10. Config (AWS Config):\n"
            "- Description: AWS Config provides configuration tracking and compliance auditing for AWS resources, helping organizations monitor "
            "changes and ensure adherence to security standards.\n"
            "- Key Findings: Ensure AWS Config is enabled to track changes across resources.\n"
            "- Recommendations: Use AWS Config to continuously monitor compliance and ensure configurations meet organizational security policies.\n\n"
        )
        
        self.document.add_paragraph(synopsis_text)

        # Add Conclusion section
        self.document.add_page_break()
        conclusion_section = self.document.add_heading("Conclusion", level=2)
        conclusion_section.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        conclusion_text = (
            "The findings from this report highlight areas of improvement across various AWS services and resources. By implementing the recommended "
            "best practices and security measures, organizations can strengthen their security posture, ensure compliance, and reduce the risk of potential "
            "security breaches. Regular monitoring, audits, and updates to security policies and configurations are essential to maintaining a secure and "
            "compliant AWS environment.\n\n"

            "Additionally, it is important to consider other critical AWS services for security and compliance purposes, such as:\n"
            "- AWS WAF (Web Application Firewall): Protects applications from common web exploits.\n"
            "- AWS Shield: Provides protection against DDoS attacks.\n"
            "- AWS Macie: Helps with data privacy by automatically discovering and classifying sensitive data.\n"
            "- AWS Security Hub: Centralizes security findings from multiple AWS services for easier management.\n"
            "- AWS Inspector: Automates security assessments and identifies vulnerabilities in EC2 instances.\n"
            "- AWS Trusted Advisor: Provides real-time guidance to help provision resources following best practices.\n\n"
            "By integrating these services into the overall security strategy, organizations can ensure a more holistic, robust approach to "
            "securing their AWS environments."
        )
        self.document.add_paragraph(conclusion_text)

        # Save the document
        output_filename = f"{os.path.splitext(self.excel_file)[0]}_report.docx"
        self.document.save(output_filename)
        print(f"Report generated: {output_filename}")
    
    def _extract_table_data(self, sheet_name):
        """
        Extract table data from a specific sheet
        
        Args:
            sheet_name (str): Name of the Excel sheet
        
        Returns:
            list: Table data
        """
        sheet = self.workbook[sheet_name]
        data = []
        
        for row in sheet.iter_rows(values_only=True):
            # Skip empty rows
            if not any(cell for cell in row):
                continue
            data.append(row)
        
        return data

def main():
    excel_file = input("Enter the path to the Excel compliance report: ").strip()
    client_name = input("Enter the client name: ").strip()
    services_link = input("Enter the link to the detailed services Excel: ").strip()
    logo_path = "/home/rajath.h@optit.india/Documents/CSPM/imp_program/opt_it_technologies_i_pvt__ltd_logo.jpeg"

    try:
        generator = ComplianceReportDocumentGenerator(excel_file, client_name, services_link, logo_path)
        generator.generate_comprehensive_report()
        
    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
